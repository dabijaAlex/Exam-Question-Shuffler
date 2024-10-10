# Copyright 2024 Dabija Alexandru

# Permission is hereby granted, free of charge, to any person obtaining a copy of this software and associated 
# documentation files (the “Software”), to deal in the Software without restriction, including without limitation 
# the rights to use, copy, modify, merge, publish, distribute, sublicense, and/or sell copies of the Software, and 
# to permit persons to whom the Software is furnished to do so, subject to the following conditions:

# The above copyright notice and this permission notice shall be included in all copies 
# or substantial portions of the Software.

# THE SOFTWARE IS PROVIDED “AS IS”, WITHOUT WARRANTY OF ANY KIND, EXPRESS OR IMPLIED, INCLUDING BUT NOT LIMITED TO 
# THE WARRANTIES OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE AUTHORS 
# OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, 
# ARISING FROM, OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE SOFTWARE.


import random
import os
from docx import Document
import shutil
from docx.shared import Pt
import re
from io import BytesIO
# from docxcompose.processor import Processor


class Create_All_files():
    
    @staticmethod
    def take_input_from_file(input_list, nr_questions_total = 9):

        partition = nr_questions_total // len(input_list)
        exercitii = []
        numbers = "0123456789"
        index_input = 0

        for i in range(0, len(input_list)):
            index_exercitii = -1
            exercitii.append([])
            exercitii[i] = []
            with open(input_list[index_input], 'r', encoding='utf-8') as f:
                for x in f:
                    if x[0] in numbers:
                        words = re.split(r'[ \t]+', x)
                        x = " ".join(words[1:])
                        exercitii[i].append(x)
                        index_exercitii += 1
                        exercitii[i][index_exercitii] = "\t" + exercitii[i][index_exercitii]
                    else:
                        # exercitii[i][index_exercitii] += "\t"
                        exercitii[i][index_exercitii] += str(x)
                for j, x in enumerate(exercitii[i]):
                    exercitii[i][j] = exercitii[i][j][:-1]


            random.shuffle(exercitii[i])
            index_input += 1
        return exercitii

    @staticmethod
    def create_Dir(parent_dir):
        directory = "SubiecteWord"
        
        path = os.path.join(parent_dir, directory) 
        
        try:
            os.mkdir(path) 
        except FileExistsError:
            a = 1
            shutil.rmtree(path)
            os.mkdir(path)
    

    @staticmethod
    def take_nr_of_elems_from_index(exercitii, nr, index):
        set = []
        nr = int(nr)
        for i in range(0, nr):
            try:
                set.append(exercitii[index + i])
            except IndexError:
                return []
        return set

    @staticmethod
    def cat_questions(bold, size, name, path_to_folder, exercitii, input_list_numbers = [], header_path = "./test.docx", batch = 1):
        index_lists = []
        for i in range(0, len(input_list_numbers)):
            index_lists.append(0)
        increment_docx = 1
        ok = 1
        while True:
            index_exercitiu = 1
            file_name = (f"{path_to_folder}/SubiecteWord/SubiectNr_v{batch}_({increment_docx:03})_.docx")
            increment_docx += 1

            index_input = 0
            cat_set = []

            for x in exercitii:
                nr = input_list_numbers[index_input]
                set = Create_All_files.take_nr_of_elems_from_index(x, nr, index_lists[index_input])
                if set == []:
                    ok = 0
                    break
                cat_set += set
                index_lists[index_input] += nr
                index_input += 1

            if ok == 0:
                break


            document = Document(header_path)
            document.save(file_name)

            document = Document(file_name)

            for x in cat_set:
                p = document.add_paragraph()
                run = p.add_run(f"{index_exercitiu}.{x}")
                run.bold = bold
                run.font.name = name
                run.font.size = Pt(size)
                index_exercitiu += 1
            document.save(file_name)


    @staticmethod
    def combine_word_documents(path_to_folder, files, merged_document):
        merged_document.add_page_break()

        for index, file in enumerate(files):
            sub_doc = Document(file)

            if index < len(files)-1:
                sub_doc.add_page_break()


            for element in sub_doc.element.body:
                merged_document.element.body.append(element)

        merged_document.save(f"{path_to_folder}/SubiecteWord/merged.docx")


    @staticmethod
    def exec_all(path_to_folder, input_list = [], input_list_numbers = [], header_path = "test.docx",
                 bold = False, size = 8, name = "Calibri", nr_of_total_files = 40):

        if path_to_folder == "":
            return False
        try:
            Create_All_files.create_Dir(path_to_folder)
        except FileNotFoundError:
            return False
        
        try:
            exercitii = Create_All_files.take_input_from_file(input_list)
        except FileNotFoundError:
            return False
        
        batch = 1
        z = 1
        while z == 1:
            Create_All_files.cat_questions(bold, size, name, path_to_folder, exercitii, input_list_numbers, header_path, batch)
            for lista_exercitii in exercitii:
                random.shuffle(lista_exercitii)
            lista = os.listdir(f"{path_to_folder}/SubiecteWord")
            print(lista)
            print(len(lista))
            if len(lista) > nr_of_total_files:
                break
            if len(lista) == nr_of_total_files:
                break
            batch += 1

        list = os.listdir(f"{path_to_folder}/SubiecteWord")
        # print(list)
        for index, x in enumerate(list):
            x = f"{path_to_folder}/SubiecteWord/{x}"
            list[index] = x

        doc1 = Document(list[0])
        list = list[1:]
        Create_All_files.combine_word_documents(path_to_folder, list, doc1)

        return True
